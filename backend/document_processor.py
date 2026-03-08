# backend/document_processor.py
import io
from typing import List, Dict

import docx


# =====================================================================
#                        CONFIGURAÇÕES GERAIS
# =====================================================================

REQUIRED_TAGS: List[str] = [
    "[NOME DO COLABORADOR]",
    "[NÚMERO DA CARTEIRA]",
    "[DIA, MÊS E ANO]",
]

# =====================================================================
#                        VALIDAÇÃO E PROCESSAMENTO EM RAM
# =====================================================================


def check_tags_in_docx(file_stream: io.BytesIO) -> List[str]:
    """
    Lê um fluxo binário de arquivo .docx direto da RAM e verifica as tags.
    """
    doc = docx.Document(file_stream)
    full_text: List[str] = []

    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    full_text.append(cell.text)

    full_text_str = " ".join(full_text)

    missing_tags = [tag for tag in REQUIRED_TAGS if tag not in full_text_str]
    return missing_tags


def fill_contract(file_stream: io.BytesIO, context: Dict[str, str]) -> io.BytesIO:
    """
    Abre o modelo binário, substitui as chaves e devolve um novo fluxo de memória.
    NENHUM arquivo é salvo no disco rígido.
    """
    doc = docx.Document(file_stream)

    # Coleta todos os parágrafos do documento, incluindo os de tabelas
    all_paragraphs = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_paragraphs.extend(cell.paragraphs)

    for para in all_paragraphs:
        for key, value in context.items():
            # A substituição pode alterar a estrutura dos runs, então usamos um loop
            # para garantir que todas as ocorrências da mesma chave sejam substituídas.
            while key in para.text:
                if not _replace_key_in_paragraph(para, key, value):
                    # Se a substituição falhar, quebra o loop para evitar repetição infinita.
                    break

    # Cria um novo arquivo binário na memória RAM para devolver ao usuário
    output_stream = io.BytesIO()
    doc.save(output_stream)
    output_stream.seek(0)

    return output_stream


def _replace_key_in_paragraph(paragraph, key, value) -> bool:
    """
    Encontra e substitui a primeira ocorrência de uma `key` em um `paragraph`,
    mesmo que a chave esteja dividida entre múltiplos `runs`.
    Retorna True se a substituição foi feita, False caso contrário.
    """
    # 1. Encontra a posição da chave no texto concatenado do parágrafo.
    text_runs = [r.text for r in paragraph.runs]
    full_text = "".join(text_runs)
    if key not in full_text:
        return False

    key_start_pos = full_text.find(key)
    key_end_pos = key_start_pos + len(key)

    # 2. Identifica os runs que contêm a chave.
    runs_to_process = []
    run_start = 0
    for i, run_text in enumerate(text_runs):
        run_end = run_start + len(run_text)
        if run_start < key_end_pos and run_end > key_start_pos:
            runs_to_process.append(i)
        run_start = run_end

    if not runs_to_process:
        return False

    # 3. Realiza a substituição mantendo a formatação.
    run_start_pos = sum(len(text) for text in text_runs[: runs_to_process[0]])
    before_text = paragraph.runs[runs_to_process[0]].text[
        : key_start_pos - run_start_pos
    ]
    last_run_start_pos = sum(len(text) for text in text_runs[: runs_to_process[-1]])
    after_text = paragraph.runs[runs_to_process[-1]].text[
        key_end_pos - last_run_start_pos :
    ]

    paragraph.runs[runs_to_process[0]].text = before_text + value
    for i in range(1, len(runs_to_process) - 1):
        paragraph.runs[runs_to_process[i]].text = ""
    if len(runs_to_process) > 1:
        paragraph.runs[runs_to_process[-1]].text = after_text
    else:
        paragraph.runs[runs_to_process[0]].text += after_text

    return True
